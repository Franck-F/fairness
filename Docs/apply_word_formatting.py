"""
Post-traitement du DOCX généré par pandoc pour le rendre conforme aux
exigences typographiques du syllabus Epitech (page 15) :

- Police : Times New Roman, 12 pt pour le texte principal, 10 pt pour les
  notes de bas de page et légendes.
- Interligne : 1.5 pour le corps, simple pour notes et citations.
- Justification : texte principal justifié.
- Pagination : numérotation des pages dans le pied de page.

Usage : python Docs/apply_word_formatting.py
Entrée/sortie : Docs/memoire_final.docx (modifié en place)
"""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

DOCX = Path(__file__).parent / "memoire_final.docx"

BODY_FONT = "Times New Roman"
BODY_SIZE = Pt(12)
NOTE_SIZE = Pt(10)
LINE_SPACING = 1.5


def _set_run_font(run, name: str, size: Pt) -> None:
    run.font.name = name
    run.font.size = size
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), name)


def _set_default_style(document: Document) -> None:
    """Force la police, taille et interligne sur le style Normal."""
    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = BODY_SIZE
    pf = normal.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), BODY_FONT)


def _format_all_paragraphs(document: Document) -> int:
    """Applique police + interligne + justification sur chaque paragraphe."""
    count = 0
    for para in document.paragraphs:
        style_name = para.style.name if para.style else ""
        is_heading = style_name.lower().startswith("heading") or "title" in style_name.lower()
        size = BODY_SIZE
        if is_heading:
            size = BODY_SIZE  # on garde 12pt mais Word boldifiera via style heading
        for run in para.runs:
            _set_run_font(run, BODY_FONT, size)
        if not is_heading:
            para.paragraph_format.line_spacing = LINE_SPACING
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            if para.paragraph_format.alignment is None:
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        count += 1
    return count


def _format_tables(document: Document) -> None:
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        _set_run_font(run, BODY_FONT, BODY_SIZE)


def _add_page_numbers(document: Document) -> None:
    """Ajoute un champ PAGE dans le pied de page de chaque section."""
    for section in document.sections:
        footer = section.footer
        if footer.paragraphs and footer.paragraphs[0].text.strip():
            para = footer.paragraphs[0]
            for run in para.runs:
                run.text = ""
        else:
            para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        # Champ PAGE
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = "PAGE \\* MERGEFORMAT"
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run._element.append(fldChar1)
        run._element.append(instrText)
        run._element.append(fldChar2)
        _set_run_font(run, BODY_FONT, NOTE_SIZE)


def _format_footnotes(document: Document) -> None:
    """Applique 10 pt aux notes de bas de page."""
    fn_part = document.part.package.parts
    for part in fn_part:
        if part.partname.endswith("/footnotes.xml"):
            xml = part.element
            for r in xml.iter(qn("w:r")):
                rPr = r.find(qn("w:rPr"))
                if rPr is None:
                    rPr = OxmlElement("w:rPr")
                    r.insert(0, rPr)
                # Set font + size
                rFonts = rPr.find(qn("w:rFonts"))
                if rFonts is None:
                    rFonts = OxmlElement("w:rFonts")
                    rPr.append(rFonts)
                for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
                    rFonts.set(qn(attr), BODY_FONT)
                sz = rPr.find(qn("w:sz"))
                if sz is None:
                    sz = OxmlElement("w:sz")
                    rPr.append(sz)
                sz.set(qn("w:val"), "20")  # 10 pt = 20 demi-points


def main() -> None:
    if not DOCX.exists():
        raise SystemExit(f"DOCX manquant : {DOCX}")

    print(f"Lecture : {DOCX.name}")
    doc = Document(str(DOCX))

    print("1/4 Style Normal : Times New Roman 12pt, interligne 1.5, justifie")
    _set_default_style(doc)

    print("2/4 Application sur tous les paragraphes")
    n = _format_all_paragraphs(doc)
    print(f"      -> {n} paragraphes formates")

    print("3/4 Tables et pied de page (numerotation)")
    _format_tables(doc)
    _add_page_numbers(doc)

    print("4/4 Notes de bas de page (Times New Roman 10pt)")
    try:
        _format_footnotes(doc)
    except Exception as exc:
        print(f"      AVERTISSEMENT : {exc}")

    doc.save(str(DOCX))
    size_kb = DOCX.stat().st_size // 1024
    print(f"\nOK : {DOCX.name} ({size_kb} K) mis a jour.")


if __name__ == "__main__":
    main()
